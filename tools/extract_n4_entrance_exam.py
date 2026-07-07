# -*- coding: utf-8 -*-
"""Extract 5 mã đề Test Đầu Vào N4 (docx) -> 1 JSON chuẩn hoá (giống schema N3).
Nguồn: KOERU_5_De_Dau_Vao_N4_N5_Level/Test đầu vào N4/*.docx
Chạy: python tools/extract_n4_entrance_exam.py
"""
import docx, re, json, io, os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.join(ROOT, "KOERU_5_De_Dau_Vao_N4_N5_Level", "Test đầu vào N4")
OUT_JSON = os.path.join(SRC_DIR, "SOURCE_OF_TRUTH_exam_data_n4.json")
OUT_JS = os.path.join(ROOT, "test-dau-vao-n4-data.js")

ANSWER_KEY_FILE = os.path.join(SRC_DIR, "KOERU_Dau_Vao_N4_Nen_Tang_N5_Dap_An_Huong_Dan_Giao_Vien_Reading_QA_Final.docx")
SET_FILES = [os.path.join(SRC_DIR, "KOERU_Dau_Vao_N4_Nen_Tang_N5_Ma_%02d_Khong_Nghe_Reading_QA_Final.docx" % i) for i in range(1, 6)]

QNUM_RE = re.compile(r'^(\d{1,2})\.\s*(.*)', re.DOTALL)


def extract_set(path):
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    prompts = {}
    for p in paras:
        m = QNUM_RE.match(p)
        if m:
            n = int(m.group(1))
            if 1 <= n <= 40 and n not in prompts:
                prompts[n] = m.group(2).strip()
    # tables[0] = header info, tables[1..40] = 4 phương án mỗi câu (2x2), tables[41] = phiếu trả lời trống
    options = {}
    for i in range(1, 41):
        t = d.tables[i]
        opts = []
        for row in t.rows:
            for cell in row.cells:
                opts.append(cell.text.strip())
        # bỏ số thứ tự "1. " ở đầu mỗi phương án
        cleaned = [re.sub(r'^\d\.\s*', '', o) for o in opts]
        options[i] = cleaned
    return prompts, options


def extract_answer_key(path):
    d = docx.Document(path)
    # table index 2..6 tương ứng mã đề 01..05 (0=info thí sinh, 1=thang điểm)
    result = {}
    for set_idx, tbl_idx in enumerate(range(2, 7), start=1):
        t = d.tables[tbl_idx]
        rows = [[c.text.strip() for c in r.cells] for r in t.rows[1:]]  # bỏ header
        qa = {}
        for r in rows:
            num, ans, correct_text, explanation, skill = r[0], r[1], r[2], r[3], r[4]
            qa[int(num)] = {
                "answer": int(ans),
                "correct_text": correct_text,
                "explanation": explanation,
                "skill": skill,
            }
        result[set_idx] = qa
    return result


answer_keys = extract_answer_key(ANSWER_KEY_FILE)

sets_out = []
for set_idx, path in enumerate(SET_FILES, start=1):
    prompts, options = extract_set(path)
    qa = answer_keys[set_idx]
    questions = []
    for n in range(1, 41):
        questions.append({
            "number": n,
            "prompt": prompts[n],
            "options": options[n],
            "answer": qa[n]["answer"],
            "correct_text": qa[n]["correct_text"],
            "explanation": qa[n]["explanation"],
            "skill": qa[n]["skill"],
        })
    sets_out.append(questions)

data = {
    "title": "KOERU – Đầu vào N4, nền tảng N5",
    "audience": "Đã hoàn thành N5, chuẩn bị học N4",
    "version": "v1-docx-extract-review-required",
    "duration_minutes": 65,
    "listening_separate": True,
    "status": "REVIEW_REQUIRED",
    "sets": sets_out,
}

with io.open(OUT_JSON, "w", encoding="utf-8") as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

with io.open(OUT_JS, "w", encoding="utf-8") as f:
    f.write("// AUTO-GENERATED từ docx (extract_n4_entrance_exam.py) — REVIEW_REQUIRED, chưa duyệt.\n")
    f.write("window.EXAM_N4 = ")
    json.dump(data, f, ensure_ascii=False, separators=(",", ":"))
    f.write(";\n")

print("OK ->", OUT_JSON)
print("OK ->", OUT_JS)
print("Sets:", len(sets_out), "| câu/mã đề:", [len(s) for s in sets_out])
