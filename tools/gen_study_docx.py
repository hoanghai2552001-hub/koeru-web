"""
KOERU — JLPT Study DOCX Generator (N5–N1)
Tạo file DOCX cho từng cấp độ JLPT với:
  - Ảnh hoạt họa stroke order (animated GIF) từ KanjiVG
  - Dữ liệu: Hán Việt, nghĩa VI, On/Kun, gợi nhớ, từ mẫu
Usage:
  python tools/gen_study_docx.py          # tất cả levels
  python tools/gen_study_docx.py N5       # chỉ N5
"""

import sys, re, os, time, json, io, urllib.request
import xml.etree.ElementTree as ET
from pathlib import Path
from PIL import Image
import aggdraw
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
ROOT       = SCRIPT_DIR.parent
JS_DIR     = ROOT / "js"
OUT_DIR    = ROOT / "output"
PNG_CACHE  = OUT_DIR / "kanji_png"
OUT_DIR.mkdir(exist_ok=True)
PNG_CACHE.mkdir(exist_ok=True)

# ── Config ────────────────────────────────────────────────────────────────────
LEVELS = ["N5", "N4", "N3", "N2", "N1"]

LEVEL_COLORS = {
    "N5": RGBColor(0x22, 0xC5, 0x5E),   # green
    "N4": RGBColor(0x3B, 0x82, 0xF6),   # blue
    "N3": RGBColor(0xF5, 0x9E, 0x0B),   # amber
    "N2": RGBColor(0xA8, 0x55, 0xF7),   # purple
    "N1": RGBColor(0xEF, 0x44, 0x44),   # red
}

LEVEL_COLORS_HEX = {
    "N5": "22C55E", "N4": "3B82F6", "N3": "F59E0B",
    "N2": "A855F7", "N1": "EF4444",
}

KANJIVG_BASE = "https://raw.githubusercontent.com/KanjiVG/kanjivg/master/kanji/{}.svg"
SVG_SIZE     = 109   # KanjiVG viewBox is 0 0 109 109
GIF_SIZE     = 160   # px per frame
GIF_DELAY    = 40    # centiseconds per frame (0.4s)
GIF_LAST     = 200   # last frame holds longer (2s)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. PARSE KANJI DATA
# ═══════════════════════════════════════════════════════════════════════════════

def parse_level_js(level: str) -> list[dict]:
    """Parse js/kanji-data-{level}.js → list of kanji dicts."""
    fpath = JS_DIR / f"kanji-data-{level.lower()}.js"
    content = fpath.read_text(encoding="utf-8")

    # Strip JS wrapper → find the array [...]
    m = re.search(r'=\s*(\[[\s\S]*\])\s*;?\s*$', content, re.MULTILINE)
    if not m:
        raise ValueError(f"Cannot find array in {fpath}")

    # Convert JS literal → JSON (minimal transform)
    txt = m.group(1)
    # quote unquoted keys
    txt = re.sub(r'([{,]\s*)([a-zA-Z_]\w*)\s*:', r'\1"\2":', txt)
    # remove trailing commas
    txt = re.sub(r',\s*([}\]])', r'\1', txt)
    # JS single-quote strings → double-quote
    txt = re.sub(r"'([^'\\]*)'", lambda m2: '"' + m2.group(1).replace('"', '\\"') + '"', txt)

    try:
        return json.loads(txt)
    except json.JSONDecodeError:
        # Fallback: brace-depth parser
        return _brace_parse(content)


def _brace_parse(content: str) -> list[dict]:
    """Extract each kanji entry using brace-depth tracking."""
    entries = []
    depth = 0; start = None
    for i, ch in enumerate(content):
        if ch == '{':
            if depth == 0: start = i
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0 and start is not None:
                block = content[start:i+1]
                entry = {}
                for key, pat in [
                    ("kanji",      r'kanji\s*:\s*"([^"]*)"'),
                    ("hanviet",    r'hanviet\s*:\s*"([^"]*)"'),
                    ("on",         r'"on"\s*:\s*"([^"]*)"'),
                    ("kun",        r'kun\s*:\s*"([^"]*)"'),
                    ("meaning",    r'meaning\s*:\s*"([^"]*)"'),
                    ("level",      r'level\s*:\s*"([^"]*)"'),
                    ("mnemonic",   r'mn_vi\s*:\s*"([^"]*)"'),
                    ("radical",    r'radical\s*:\s*"([^"]*)"'),
                ]:
                    mo = re.search(pat, block)
                    if mo: entry[key] = mo.group(1)
                for key in ("stroke", "freq_rank", "grade"):
                    mo = re.search(rf'\b{key}\s*:\s*(\d+)', block)
                    if mo: entry[key] = int(mo.group(1))
                # words array
                words = re.findall(r'\{"w"\s*:\s*"([^"]*)","r"\s*:\s*"([^"]*)","m"\s*:\s*"([^"]*)"\}', block)
                entry["words"] = [{"w": w, "r": r, "m": mm} for w, r, mm in words]
                if entry.get("kanji"): entries.append(entry)
                start = None
    return entries


# ═══════════════════════════════════════════════════════════════════════════════
# 2. ANIMATED GIF FROM KANJIVG
# ═══════════════════════════════════════════════════════════════════════════════

def kanji_to_hex(ch: str) -> str:
    return f"{ord(ch):05x}"


def fetch_svg(kanji: str) -> str | None:
    """Fetch KanjiVG SVG text, return None on error."""
    hex_code = kanji_to_hex(kanji)
    url = KANJIVG_BASE.format(hex_code)
    try:
        with urllib.request.urlopen(url, timeout=10) as r:
            return r.read().decode("utf-8")
    except Exception:
        return None


def svg_to_paths(svg_text: str) -> list[str]:
    """Extract ordered stroke path 'd' attributes from KanjiVG SVG."""
    root = ET.fromstring(svg_text)
    paths = root.findall(".//{http://www.w3.org/2000/svg}path")
    # Filter only stroke paths (id contains '-s')
    stroke_paths = [p for p in paths if p.get("id", "").find("-s") != -1]
    # Fallback: use all paths
    if not stroke_paths:
        stroke_paths = paths
    return [p.get("d", "") for p in stroke_paths if p.get("d")]


def render_frame(active_ds: list[str], all_ds: list[str], size: int = GIF_SIZE,
                 stroke_num: int = 0) -> Image.Image:
    """Render one animation frame: active strokes black, rest light gray."""
    img = Image.new("RGB", (size, size), (255, 255, 255))
    canvas = aggdraw.Draw(img)
    scale = size / SVG_SIZE

    # Guide strokes (very light)
    pen_guide = aggdraw.Pen((210, 210, 210), width=max(1, 2.2 * scale))
    for d in all_ds:
        try: canvas.symbol((0, 0), aggdraw.Symbol(d), pen_guide)
        except: pass

    # Previous strokes (medium gray)
    pen_done = aggdraw.Pen((90, 90, 90), width=max(1, 3 * scale))
    for d in active_ds[:-1]:
        try: canvas.symbol((0, 0), aggdraw.Symbol(d), pen_done)
        except: pass

    # Current stroke (accent color, thicker)
    pen_current = aggdraw.Pen((220, 50, 50), width=max(1, 4 * scale))
    if active_ds:
        try: canvas.symbol((0, 0), aggdraw.Symbol(active_ds[-1]), pen_current)
        except: pass

    canvas.flush()

    # Draw stroke number badge
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    badge_r = int(10 * scale)
    badge_x = size - badge_r - 4
    badge_y = 4
    draw.ellipse(
        [badge_x - badge_r, badge_y, badge_x + badge_r, badge_y + badge_r * 2],
        fill=(220, 50, 50)
    )
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", int(9 * scale))
    except:
        font = ImageFont.load_default()
    label = str(stroke_num)
    bbox = draw.textbbox((0, 0), label, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(
        (badge_x - tw // 2, badge_y + badge_r - th // 2),
        label, fill=(255, 255, 255), font=font
    )
    return img


def make_stroke_order_image(kanji: str, svg_text: str) -> Image.Image:
    """Create a static stroke order strip image (N frames side by side)."""
    ds = svg_to_paths(svg_text)
    if not ds:
        return None
    n = len(ds)
    FRAME = GIF_SIZE
    img_strip = Image.new("RGB", (FRAME * n, FRAME), (255, 255, 255))
    for i in range(n):
        frame = render_frame(ds[:i+1], ds, FRAME, stroke_num=i+1)
        img_strip.paste(frame, (FRAME * i, 0))
    return img_strip


def make_animated_gif(kanji: str, svg_text: str) -> bytes:
    """Create animated GIF bytes from KanjiVG SVG."""
    ds = svg_to_paths(svg_text)
    if not ds:
        return None
    frames = []
    durations = []
    for i in range(len(ds)):
        frame = render_frame(ds[:i+1], ds, GIF_SIZE, stroke_num=i+1)
        frames.append(frame.convert("P", palette=Image.ADAPTIVE, colors=64))
        durations.append(GIF_DELAY * 10)  # PIL uses milliseconds
    durations[-1] = GIF_LAST * 10

    buf = io.BytesIO()
    frames[0].save(
        buf, format="GIF",
        save_all=True,
        append_images=frames[1:],
        duration=durations,
        loop=0,
        optimize=False,
    )
    return buf.getvalue()


def get_or_create_gif(kanji: str) -> bytes | None:
    """Get cached GIF or create new one."""
    hex_code = kanji_to_hex(kanji)
    cache_path = PNG_CACHE / f"{hex_code}.gif"
    if cache_path.exists():
        return cache_path.read_bytes()

    svg_text = fetch_svg(kanji)
    if not svg_text:
        print(f"    ⚠  No SVG for {kanji}")
        return None

    gif_bytes = make_animated_gif(kanji, svg_text)
    if gif_bytes:
        cache_path.write_bytes(gif_bytes)
    return gif_bytes


# ═══════════════════════════════════════════════════════════════════════════════
# 3. DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top="none", bottom="none", left="none", right="none",
                     color="E2E8F0", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        if val != "none":
            el.set(qn("w:color"), color)
            el.set(qn("w:sz"), sz)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_col_widths(table, widths_cm: list[float]):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(widths_cm[i] * 567)))  # 1cm = 567 twips
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def para_add_run(para, text: str, bold=False, italic=False,
                 size_pt: int = 10, color: RGBColor = None, font_name: str = "Noto Sans JP"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    try:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass
    return run


def cell_para(cell, text: str, bold=False, size_pt=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    """Clear cell and add a paragraph."""
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        para_add_run(para, text, bold=bold, size_pt=size_pt, color=color)
    return para


# ═══════════════════════════════════════════════════════════════════════════════
# 4. BUILD DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def add_cover_page(doc: Document, level: str, kanji_list: list[dict]):
    color = LEVEL_COLORS[level]
    hex_c = LEVEL_COLORS_HEX[level]

    # Big level title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(f"KOERU JLPT {level}")
    run.bold = True
    run.font.size = Pt(40)
    run.font.color.rgb = color

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after = Pt(12)
    run2 = sub_para.add_run("Kanji Study Book — Bảng Học Kanji")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Stats
    total_words = sum(len(k.get("words", [])) for k in kanji_list)
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_para.paragraph_format.space_before = Pt(8)
    stats_para.paragraph_format.space_after = Pt(24)
    run3 = stats_para.add_run(
        f"{len(kanji_list)} Kanji   ·   ~{total_words} Từ mẫu   ·   Stroke order GIF"
    )
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # How to use
    how_table = doc.add_table(rows=1, cols=1)
    how_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = how_table.rows[0].cells[0]
    set_cell_bg(cell, "F8FAFC")
    set_cell_borders(cell, top="single", bottom="single", left="single", right="single",
                     color="CBD5E1", sz="4")
    cell.width = Cm(14)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Cách dùng:")
    run.bold = True; run.font.size = Pt(11)
    for line in [
        "• GIF hoạt họa → theo dõi thứ tự nét viết từng bước",
        "• Cột 'Gợi nhớ' → câu chuyện kết nối hình ảnh với nghĩa",
        "• Cột 'Từ mẫu' → ứng dụng ngay vào từ vựng thực tế",
        "• In hai mặt, gấp đôi: nhìn Kanji → đoán nghĩa → lật kiểm tra",
    ]:
        lp = cell.add_paragraph(line)
        lp.paragraph_format.space_before = Pt(1)
        lp.paragraph_format.space_after = Pt(1)
        lp.runs[0].font.size = Pt(10)
        lp.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    end_p = cell.add_paragraph()
    end_p.paragraph_format.space_before = Pt(4)
    end_p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def build_kanji_table(doc: Document, kanji_list: list[dict], level: str):
    """Build the main kanji table with GIF images."""
    color = LEVEL_COLORS[level]
    hex_c  = LEVEL_COLORS_HEX[level]

    # Section header
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_before = Pt(0)
    hdr.paragraph_format.space_after = Pt(6)
    run = hdr.add_run(f"Bang Kanji {level}  ({len(kanji_list)} chu)")
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = color

    # Table: GIF | Kanji+Han Viet | On/Kun | Nghia + Goi nho | Tu mau
    # Widths: 4cm | 3cm | 3cm | 5.5cm | 5.5cm  = 21cm total
    COL_W = [4.0, 2.8, 3.2, 5.0, 6.0]

    # Header row
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    hdr_row = tbl.rows[0]
    hdr_labels = ["Stroke Order", "Kanji / Han Viet", "On / Kun", "Nghia & Goi nho", "Tu mau (vi du)"]
    for i, (cell, lbl) in enumerate(zip(hdr_row.cells, hdr_labels)):
        set_cell_bg(cell, hex_c)
        cell_para(cell, lbl, bold=True, size_pt=9,
                  color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=4, space_after=4)

    total = len(kanji_list)
    for idx, entry in enumerate(kanji_list):
        kanji_char = entry.get("kanji", "?")
        print(f"  [{idx+1}/{total}] {kanji_char}", end=" ", flush=True)

        row_color = "FFFFFF" if idx % 2 == 0 else "F8FAFC"
        row = tbl.add_row()

        # ── Col 0: Animated GIF ──────────────────────────────────────────────
        gif_cell = row.cells[0]
        set_cell_bg(gif_cell, row_color)
        gif_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        gif_bytes = get_or_create_gif(kanji_char)
        if gif_bytes:
            print("✓", end=" ", flush=True)
            gif_para = gif_cell.paragraphs[0]
            gif_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            gif_para.paragraph_format.space_before = Pt(3)
            gif_para.paragraph_format.space_after  = Pt(3)
            gif_run = gif_para.add_run()
            gif_run.add_picture(io.BytesIO(gif_bytes), width=Cm(3.2))
        else:
            print("✗", end=" ", flush=True)
            # Fallback: show kanji character large
            p = gif_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(kanji_char)
            r.font.size = Pt(40)
            r.font.color.rgb = color

        # ── Col 1: Kanji char + Han Viet + strokes ───────────────────────────
        info_cell = row.cells[1]
        set_cell_bg(info_cell, row_color)
        info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Big kanji
        kp = info_cell.paragraphs[0]
        kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kp.paragraph_format.space_before = Pt(4)
        kp.paragraph_format.space_after  = Pt(0)
        kr = kp.add_run(kanji_char)
        kr.bold = True
        kr.font.size = Pt(32)
        kr.font.color.rgb = color
        try:
            kr.font.name = "Noto Sans JP"
            kr._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans JP")
        except: pass

        # Han Viet
        hv = info_cell.add_paragraph()
        hv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hv.paragraph_format.space_before = Pt(1)
        hv.paragraph_format.space_after  = Pt(1)
        hvr = hv.add_run(entry.get("hanviet", ""))
        hvr.bold = True; hvr.font.size = Pt(10)
        hvr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Stroke count
        sc = info_cell.add_paragraph()
        sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc.paragraph_format.space_before = Pt(0)
        sc.paragraph_format.space_after  = Pt(4)
        scr = sc.add_run(f"{entry.get('stroke', '?')} net")
        scr.font.size = Pt(8)
        scr.font.color.rgb = RGBColor(0xAB, 0xB5, 0xBE)

        # ── Col 2: On / Kun ──────────────────────────────────────────────────
        read_cell = row.cells[2]
        set_cell_bg(read_cell, row_color)
        read_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # On reading
        on_p = read_cell.paragraphs[0]
        on_p.paragraph_format.space_before = Pt(4)
        on_p.paragraph_format.space_after  = Pt(2)
        on_label = on_p.add_run("On: ")
        on_label.bold = True; on_label.font.size = Pt(8)
        on_label.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        on_val = on_p.add_run(entry.get("on", "—"))
        on_val.font.size = Pt(9)
        on_val.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        # Kun reading
        kun_p = read_cell.add_paragraph()
        kun_p.paragraph_format.space_before = Pt(0)
        kun_p.paragraph_format.space_after  = Pt(4)
        kun_label = kun_p.add_run("Kun: ")
        kun_label.bold = True; kun_label.font.size = Pt(8)
        kun_label.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        kun_val = kun_p.add_run(entry.get("kun", "—"))
        kun_val.font.size = Pt(9)
        kun_val.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

        # ── Col 3: Nghia + Goi nho ───────────────────────────────────────────
        mn_cell = row.cells[3]
        set_cell_bg(mn_cell, row_color)
        mn_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Meaning
        meaning = entry.get("meaning", "")
        mp = mn_cell.paragraphs[0]
        mp.paragraph_format.space_before = Pt(4)
        mp.paragraph_format.space_after  = Pt(3)
        ml = mp.add_run("Nghia: ")
        ml.bold = True; ml.font.size = Pt(8)
        ml.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        mv = mp.add_run(meaning)
        mv.bold = True; mv.font.size = Pt(9)
        mv.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Mnemonic (goi nho)
        mnemonic = re.sub(r'\s*\[\d+\]', '', entry.get("mnemonic", entry.get("mn_vi", ""))).strip()
        if mnemonic:
            mnp = mn_cell.add_paragraph()
            mnp.paragraph_format.space_before = Pt(0)
            mnp.paragraph_format.space_after  = Pt(4)
            mnl = mnp.add_run("Goi nho: ")
            mnl.bold = True; mnl.font.size = Pt(8)
            mnl.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            mnv = mnp.add_run(mnemonic[:120])
            mnv.italic = True; mnv.font.size = Pt(8)
            mnv.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # ── Col 4: Tu mau ────────────────────────────────────────────────────
        words_cell = row.cells[4]
        set_cell_bg(words_cell, row_color)
        words_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        words = entry.get("words", [])[:5]
        first = True
        for w in words:
            wp = words_cell.paragraphs[0] if first else words_cell.add_paragraph()
            first = False
            wp.paragraph_format.space_before = Pt(1)
            wp.paragraph_format.space_after  = Pt(1)
            wr1 = wp.add_run(f"{w.get('w','')} ")
            wr1.bold = True; wr1.font.size = Pt(9)
            wr1.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            try:
                wr1.font.name = "Noto Sans JP"
                wr1._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans JP")
            except: pass
            wr2 = wp.add_run(f"({w.get('r','')}) ")
            wr2.font.size = Pt(8)
            wr2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            wr3 = wp.add_run(w.get('m',''))
            wr3.italic = True; wr3.font.size = Pt(8)
            wr3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        print()  # newline after each kanji

    # Apply column widths
    set_table_col_widths(tbl, COL_W)

    doc.add_paragraph()  # spacing after table


def setup_document(level: str) -> Document:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.page_width    = Cm(29.7)   # A4 landscape
        section.page_height   = Cm(21.0)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    return doc


def add_footer(doc: Document, level: str):
    color = LEVEL_COLORS[level]
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("KOERU Japanese · ")
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor(0xAB, 0xB5, 0xBE)
        r2 = fp.add_run(f"JLPT {level} Study Book")
        r2.bold = True; r2.font.size = Pt(8)
        r2.font.color.rgb = color
        r3 = fp.add_run("  ·  koeruapp.com")
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(0xAB, 0xB5, 0xBE)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def generate_level(level: str):
    print(f"\n{'='*60}")
    print(f"  Generating KOERU_JLPT_{level}.docx ...")
    print(f"{'='*60}")

    kanji_list = parse_level_js(level)
    print(f"  Parsed {len(kanji_list)} kanji")

    doc = setup_document(level)
    add_cover_page(doc, level, kanji_list)
    build_kanji_table(doc, kanji_list, level)
    add_footer(doc, level)

    out_path = OUT_DIR / f"KOERU_JLPT_{level}.docx"
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size // 1024
    print(f"\n  ✅ Saved: {out_path.name}  ({size_kb} KB)")
    return out_path


def main():
    targets = sys.argv[1:] if len(sys.argv) > 1 else LEVELS
    targets = [t.upper() for t in targets if t.upper() in LEVELS]
    if not targets:
        print(f"Usage: python tools/gen_study_docx.py [{'|'.join(LEVELS)}]")
        sys.exit(1)

    print(f"KOERU JLPT DOCX Generator")
    print(f"Targets: {', '.join(targets)}")
    print(f"GIF cache: {PNG_CACHE}")

    for level in targets:
        t0 = time.time()
        generate_level(level)
        print(f"  Time: {time.time()-t0:.1f}s")

    print("\nDone! Files in:", OUT_DIR)


if __name__ == "__main__":
    main()
